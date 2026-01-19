
import os
import re
import json
import glob
import argparse
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime

# ========= CLI FLAGS + Interactive Filename Prompt (with last-used memory) =========
parser = argparse.ArgumentParser(description="HKPF Posting Summary Processor")
parser.add_argument("--export-unknowns", action="store_true",
                    help="Export unresolved role/location tokens to hkpf_unknowns.json")
parser.add_argument("--combine-sheets", action="store_true",
                    help="Combine all sheets (default uses first sheet only)")
parser.add_argument("--file", default=None,
                    help="Path to Excel file (.xlsx). If omitted, you will be prompted.")
args = parser.parse_args()

EXPORT_UNKNOWNS = bool(args.export_unknowns)
COMBINE_SHEETS = bool(args.combine_sheets)

HKPF_LAST_PATH = Path(".hkpf_last.json")

def list_xlsx(cwd: Path) -> list[Path]:
    return sorted((Path(p) for p in glob.glob(str(cwd / "*.xlsx"))), key=lambda p: p.name.lower())

def resolve_input_filename(user_text: str, cwd: Path) -> Path | None:
    """
    Resolve a user-entered name to an .xlsx path in the current directory.

    Accepts:
      - bare stem: 'test1' -> finds test1.xlsx (case-insensitive) in cwd
      - full filename: 'PostingSummary.xlsx'
      - relative/absolute paths with .xlsx extension
    """
    if not user_text:
        return None

    s = str(user_text).strip().strip("\"'")
    p = Path(s)

    # If user provided a filename with .xlsx extension
    if p.suffix.lower() == ".xlsx":
        if p.is_file():
            return p.resolve()
        candidate = (cwd / p).resolve()
        return candidate if candidate.is_file() else None

    # Treat it as a stem and search for *.xlsx in current folder (case-insensitive)
    stem = p.name
    candidate = cwd / f"{stem}.xlsx"
    if candidate.is_file():
        return candidate.resolve()

    for m in list_xlsx(cwd):
        if m.stem.lower() == stem.lower():
            return m.resolve()

    return None

def load_last_used() -> str | None:
    try:
        if HKPF_LAST_PATH.exists():
            with open(HKPF_LAST_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
                val = data.get("last_file")
                return val if val else None
    except Exception:
        return None
    return None

def save_last_used(path_str: str) -> None:
    try:
        with open(HKPF_LAST_PATH, "w", encoding="utf-8") as f:
            json.dump({"last_file": path_str}, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def prompt_for_file(cwd: Path) -> Path:
    """
    Prompt for a valid .xlsx filename.
    - Shows last-used file as default if available (press Enter to accept).
    - On first failure, prints a list of .xlsx files in the folder.
    """
    last_used = load_last_used()
    attempts = 3
    first_error = True

    while attempts > 0:
        prompt = "Enter Excel file name (e.g., PostingSummary or PostingSummary.xlsx)"
        if last_used and Path(last_used).is_file():
            prompt += f" [default: {Path(last_used).name}]"
        prompt += ": "

        raw = input(prompt).strip()

        # Accept default last-used if user hits Enter and we have it
        if not raw and last_used and Path(last_used).is_file():
            return Path(last_used).resolve()

        # Accept empty to try 'PostingSummary.xlsx'
        if not raw and not last_used:
            fallback = cwd / "PostingSummary.xlsx"
            if fallback.is_file():
                return fallback.resolve()
            print("[Error] Default 'PostingSummary.xlsx' not found in current folder.")
            attempts -= 1
            if first_error:
                found = list_xlsx(cwd)
                if found:
                    print("These .xlsx files are in this folder:")
                    for p in found:
                        print("  -", p.name)
                else:
                    print("(No .xlsx files found in current folder.)")
                first_error = False
            continue

        resolved = resolve_input_filename(raw, cwd)
        if resolved and resolved.is_file():
            return resolved

        print(f"[Error] Could not find an .xlsx file for: {raw!r}. "
              "Try again with a valid file that exists in this folder.")
        attempts -= 1
        if first_error:
            found = list_xlsx(cwd)
            if found:
                print("These .xlsx files are in this folder:")
                for p in found:
                    print("  -", p.name)
            else:
                print("(No .xlsx files found in current folder.)")
            first_error = False

    raise SystemExit("Exiting: no valid Excel file was provided.")

# Determine final file_path
CWD = Path(os.getcwd())
if args.file:
    fp = resolve_input_filename(args.file, CWD)
    if not fp:
        raise SystemExit(f"[Error] --file '{args.file}' not found as an .xlsx in {CWD}.")
    file_path = str(fp)
else:
    fp = prompt_for_file(CWD)
    file_path = str(fp)

save_last_used(file_path)

# ========= VOCAB FILES (learning loop) =========
VOCAB_PATH = Path("hkpf_vocab.json")
UNKNOWN_PATH = Path("hkpf_unknowns.json")

# ========= STARTER VOCAB (seeded) =========
STARTER_ROLE_EXPANSIONS = {
    # Rank-like (for detection; not printed as roles when from Post Type)
    "CSP": "Chief Superintendent of Police",
    "SSP": "Senior Superintendent of Police",
    "SP": "Superintendent of Police",
    "CIP": "Chief Inspector of Police",
    "SIP": "Senior Inspector of Police",
    "IP": "Inspector of Police",
    "PI": "Probationary Inspector of Police",
    "SSGT": "Station Sergeant",
    "SGT": "Sergeant",
    "SPC": "Senior Police Constable",
    "PC": "Police Constable",

    # District command (positions; print as roles)
    "DC": "District Commander",
    "DDC": "Deputy District Commander",
    "ADC": "Assistant District Commander",

    # Portfolios / teams / functions
    "ADM": "Administration",
    "A&S": "Administration and Support",
    "CRM": "Crime",
    "ES": "Efficiency Studies",
    "RI": "Research and Inspections",
    "CTRL": "Command and Control (Control Room)",
    "GEN": "General",
    "FLD": "Field",

    # Investigation / squads / sub‑units
    "PSU 1": "Patrol Sub-unit 1",
    "PSU 2": "Patrol Sub-unit 2",
    "PSU 3": "Patrol Sub-unit 3",
    "PSU 4": "Patrol Sub-unit 4",
    "TFSU": "Task Force Sub-unit",

    "DVIT 1": "Divisional Investigation Team 1",
    "DVIT 2": "Divisional Investigation Team 2",
    "DVIT 3": "Divisional Investigation Team 3",
    "DVIT 4": "Divisional Investigation Team 4",
    "DVIT 5": "Divisional Investigation Team 5",
    "DVIT 6": "Divisional Investigation Team 6",
    "DVIT 7": "Divisional Investigation Team 7",
    "DVIT 8": "Divisional Investigation Team 8",

    "SDS 1": "Special Duties Squad 1",
    "DSDS 2": "District Special Duties Squad 2",

    # Training / events
    "SYMPOSIUM": "Symposium",
    "RPC TRG (INTAKE)": "Recruit Police Constable Training (Intake)",

    # Seen in data
    "CS&INT": "Counterfeit, Support and Intelligence",
    "INP 2": "Inspection 2",
    "AUX": "Auxiliary",
    "SCIU": "Security Company and Guarding Services Bill-Police Inspection Team",
    "SA": "Security Advisory Section",
    "PCRO": "Police Community Relations Office",
}

STARTER_LOCATION_ALIASES = {
    # Bureaux / Formations
    "CCB": "COMMERCIAL CRIME BUREAU",
    "C DIV CCB": "COMMERCIAL CRIME BUREAU",
    "C DIVISION COMMERCIAL CRIME BUREAU": "COMMERCIAL CRIME BUREAU",
    "CPB": "CRIME PREVENTION BUREAU",
    "PPRB": "POLICE PUBLIC RELATIONS BRANCH",

    # Regions / HQs / Command centres
    "RCCC HKI": "REGIONAL COMMAND AND CONTROL CENTRE HONG KONG ISLAND",
    "HKI": "HONG KONG ISLAND REGIONAL HEADQUARTERS",
    "OPS": "OPERATIONS WING",
    "PTU": "POLICE TACTICAL UNIT",

    # Districts / Divisions (common codes)
    "CDIST": "CENTRAL DISTRICT",
    "WDIST": "WESTERN DISTRICT",
    "EDIST": "EASTERN DISTRICT",
    "WCH DIV": "WAN CHAI DIVISION",
    "WCH DIST": "WAN CHAI DISTRICT",
    "CDIV": "CENTRAL DIVISION",
    "WDIV": "WESTERN DIVISION",
    "NPDIV": "NORTH POINT DIVISION",
    "STYSDIV": "STANLEY SUB-DIVISION",
    "ABDDIV": "ABERDEEN DIVISION",

    # Training / PTU
    "PTU A": "POLICE TACTICAL UNIT (A COMPANY)",
    "PTS": "POLICE TRAINING SCHOOL",

    # Previously added
    "CAPO": "COMPLAINTS AGAINST POLICE OFFICE",
    "CAPO HKI": "COMPLAINTS AGAINST POLICE OFFICE HONG KONG ISLAND",
    "SQ": "SERVICE QUALITY WING",
    "SUPPORT": "SUPPORT WING",
    "IST": "IN-SERVICE TRAINING",
    "PC TRG": "POLICE CONSTABLE TRAINING DIVISION",
    "TRVE SUP": "TRAINING RESERVE SUPPORT WING",
    "RATU KW": "REGIONAL ANTI TRIAD UNIT KOWLOON WEST",
    "RIU KW": "REGIONAL INTELLIGENCE UNIT KOWLOON WEST",
    "RCCC NTN": "REGIONAL COMMAND AND CONTROL CENTRE NEW TERRITORIES NORTH",
    "EU NTN": "EMERGENCY UNIT NEW TERRITORIES NORTH",
    "KW": "KOWLOON WEST REGIONAL HEADQUARTERS",
    "CRM KW": "CRIME KOWLOON WEST REGIONAL HEADQUARTERS",
    "SMPDIST": "SAU MAU PING DISTRICT",
    "TWDIST": "TSUEN WAN DISTRICT",
    "MKDIST": "MONG KOK DISTRICT",
}

# ========= UTIL: Load/Save vocab and merge seeds =========
def load_vocab():
    if VOCAB_PATH.exists():
        with open(VOCAB_PATH, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = {}
    data.setdefault("version", 1)
    data.setdefault("role_expansions", {})
    data.setdefault("location_aliases", {})

    data["role_expansions"].update({k: data["role_expansions"].get(k, v)
                                    for k, v in STARTER_ROLE_EXPANSIONS.items()})
    data["location_aliases"].update({k: data["location_aliases"].get(k, v)
                                     for k, v in STARTER_LOCATION_ALIASES.items()})
    return data

def save_vocab(vocab):
    with open(VOCAB_PATH, "w", encoding="utf-8") as f:
        json.dump(vocab, f, ensure_ascii=False, indent=2)

# ========= STARTUP =========
print("Working directory:", os.getcwd())
print("File exists?", os.path.exists(file_path))

xls = pd.ExcelFile(file_path)
print("Sheets found:", xls.sheet_names)

# ========= STEP 1: READ + DETECT HEADER per sheet =========
def detect_header_row(df_raw: pd.DataFrame) -> int:
    target_headers = {'date start', 'date end', 'post type'}
    header_row = None
    for i in range(min(60, len(df_raw))):
        row_vals = df_raw.iloc[i].astype(str).str.strip().str.lower().tolist()
        if target_headers.issubset(set(row_vals)):
            header_row = i
            break
    if header_row is None:
        raise ValueError("Couldn't find header row with Date Start/Date End/Post Type. Check Excel manually.")
    return header_row

def load_sheet(sheet_name: str) -> pd.DataFrame:
    raw = pd.read_excel(file_path, sheet_name=sheet_name, header=None)
    print("\nRaw preview (first 15 rows):")
    print(raw.head(15).to_string(index=True, header=False))
    header_row = detect_header_row(raw)
    print(f"\nDetected header row at index: {header_row}")
    df_local = pd.read_excel(file_path, sheet_name=sheet_name, header=header_row)
    df_local = df_local.loc[:, ~df_local.columns.astype(str).str.startswith('Unnamed')]
    df_local = df_local.dropna(how='all')
    print("\nDetected columns after cleanup (original):")
    print(list(df_local.columns))
    return df_local

if COMBINE_SHEETS:
    frames = [load_sheet(sh) for sh in xls.sheet_names]
    df = pd.concat(frames, ignore_index=True)
else:
    df = load_sheet(xls.sheet_names[0])

# ========= STEP 2: NORMALISE COLUMN NAMES =========
def snake(s: str) -> str:
    return re.sub(r'\s+', '_', str(s).strip().lower())

aliases = {
    'date_start': 'date_start',
    'date_start_(description)': 'date_start_desc',
    'date_end': 'date_end',
    'date_end_(description)': 'date_end_desc',
    'post_type': 'post_type',
    'post_type_(description)': 'post_type_desc',
    'designation': 'designation',
    'designation_(description)': 'designation_desc',
    'location': 'location',
    'location_(description)': 'location_desc',
    'formation': 'formation',
    'major_formation': 'major_formation',
    'major_formation_(description)': 'major_formation_desc',
}
df = df.rename(columns={c: aliases.get(snake(c), snake(c)) for c in df.columns})

# Parse dates
for col in ['date_start', 'date_end']:
    if col in df.columns:
        df[col] = pd.to_datetime(df[col], errors='coerce')

# Ensure presence of text columns
for col in ['post_type', 'post_type_desc', 'designation', 'designation_desc', 'location', 'location_desc']:
    if col not in df.columns:
        df[col] = ""

# ========= STEP 3: RANK MAPPING (IP/SIP preserved & enforced) =========
rank_order = ['PC', 'SPC', 'SGT', 'SSGT', 'PI', 'IP', 'SIP', 'CIP', 'SP', 'SSP', 'CSP', 'ACP', 'SACP', 'DCP', 'CP']
rank_index = {r: i for i, r in enumerate(rank_order)}
rank_index['IP/SIP'] = rank_index['IP']  # comparison level

rank_map = {
    'pc': 'PC', 'police constable': 'PC',
    'spc': 'SPC', 'senior police constable': 'SPC',
    'sgt': 'SGT', 'sergeant': 'SGT',
    'ssgt': 'SSGT', 'station sergeant': 'SSGT',
    'pi': 'PI', 'probationary inspector': 'PI', 'probationary inspector of police': 'PI',
    'ip': 'IP', 'insp': 'IP', 'inspector': 'IP', 'inspector of police': 'IP',
    'sip': 'SIP', 'sr insp': 'SIP', 'sen insp': 'SIP', 'senior insp': 'SIP', 'senior inspector': 'SIP', 'senior inspector of police': 'SIP',
    'cip': 'CIP', 'ch insp': 'CIP', 'chief insp': 'CIP', 'chief inspector': 'CIP', 'chief inspector of police': 'CIP',
    'sp': 'SP', 'superintendent': 'SP', 'superintendent of police': 'SP',
    'ssp': 'SSP', 'senior superintendent': 'SSP', 'senior superintendent of police': 'SSP',
    'csp': 'CSP', 'chief superintendent': 'CSP', 'chief superintendent of police': 'CSP',
    'acp': 'ACP', 'assistant commissioner': 'ACP', 'assistant commissioner of police': 'ACP',
    'sacp': 'SACP', 'senior assistant commissioner': 'SACP', 'senior assistant commissioner of police': 'SACP',
    'dcp': 'DCP', 'deputy commissioner': 'DCP', 'deputy commissioner of police': 'DCP',
    'cp': 'CP', 'commissioner': 'CP', 'commissioner of police': 'CP',
}

acting_tokens_pattern = re.compile(r'\b(acting|actg|a/|ag\.|temp|temporary|acting up)\b', flags=re.IGNORECASE)

def looks_like_ip_sip(text: str) -> bool:
    s = str(text or "").lower()
    s = acting_tokens_pattern.sub('', s)
    s = s.replace('\\', '/')
    s = re.sub(r'[().,;]', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    s2 = s
    s2 = re.sub(r'\bsenior\s+inspector(?:\s+of\s+police)?\b', 'sip', s2)
    s2 = re.sub(r'\bsr\s+insp\b', 'sip', s2)
    s2 = re.sub(r'\bsen\s+insp\b', 'sip', s2)
    s2 = re.sub(r'\binspector(?:\s+of\s+police)?\b', 'ip', s2)
    s2 = re.sub(r'\binsp\b', 'ip', s2)
    tokens = set(re.split(r'[^a-z/]+', s2))
    tokens.discard('')
    return ('ip' in tokens and 'sip' in tokens) or ('ip/sip' in s2)

def map_rank(text: str):
    if looks_like_ip_sip(text):
        return 'IP/SIP'
    s = str(text or "").strip().lower()
    s = acting_tokens_pattern.sub('', s)
    s = s.replace('\\', '/')
    s = re.sub(r'[().,;]', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    if not s:
        return None
    if s in rank_map:
        v = rank_map[s]
    else:
        v = None
        for k, val in rank_map.items():
            if re.search(r'\b' + re.escape(k) + r'\b', s):
                v = val
                break
    if v in {'IP', 'SIP'}:
        return 'IP/SIP'
    return v

# ========= STEP 4: ACTING/TEMP DETECTION =========
acting_kw_pattern = re.compile(r'(?i)\b(temp|temporary|acting|actg|acting up|a/|ag\.)\b|署任|代任|代理')

def is_acting(row) -> bool:
    fields = [
        str(row.get('designation', '') or ''),
        str(row.get('designation_desc', '') or ''),
        str(row.get('post_type', '') or ''),
        str(row.get('post_type_desc', '') or ''),
    ]
    haystack = ' || '.join(fields)
    return bool(acting_kw_pattern.search(haystack))

# ========= STEP 5: SORT & PREP =========
if 'date_start' in df.columns and df['date_start'].notna().any():
    df = df.sort_values(by=['date_start', 'date_end'], ascending=[True, True], na_position='last').reset_index(drop=True)
else:
    df = df.reset_index(drop=True)

df['reported_rank'] = (df['post_type'].astype(str) + ' || ' + df.get('post_type_desc', '').astype(str)).apply(map_rank)
df['acting_flag'] = df.apply(is_acting, axis=1)

# ========= STEP 6: TRUE (SUBSTANTIVE) RANK =========
rows = df.to_dict(orient='records')

def future_has_lower_than(rows, from_idx, ref_rank):
    if ref_rank is None:
        return False
    ref_idx = rank_index.get(ref_rank, -1)
    for j in range(from_idx, len(rows)):
        if bool(rows[j].get('acting_flag')):
            continue
        rj = rows[j].get('reported_rank')
        if rj is None:
            continue
        if rank_index.get(rj, -1) < ref_idx:
            return True
    return False

current_substantive_rank = None
for i, r in enumerate(rows):
    rep = r.get('reported_rank')
    acting = bool(r.get('acting_flag'))

    if current_substantive_rank is None:
        if rep is None:
            r['true_rank'] = None
        else:
            looks_acting = acting or future_has_lower_than(rows, i + 1, rep)
            if looks_acting:
                r['true_rank'] = None
            else:
                current_substantive_rank = rep
                r['true_rank'] = current_substantive_rank
        continue

    if rep is None:
        r['true_rank'] = current_substantive_rank
        continue

    if acting:
        r['true_rank'] = current_substantive_rank
        continue

    cur_idx = rank_index.get(current_substantive_rank, -1)
    rep_idx = rank_index.get(rep, -1)

    if rep_idx > cur_idx:
        if future_has_lower_than(rows, i + 1, rep):
            r['true_rank'] = current_substantive_rank
        else:
            current_substantive_rank = rep
            r['true_rank'] = current_substantive_rank
    else:
        r['true_rank'] = current_substantive_rank

first_true = next((rr.get('true_rank') for rr in rows if rr.get('true_rank') is not None), None)
if first_true is not None:
    for r in rows:
        if r.get('true_rank') is None:
            r['true_rank'] = first_true

df = pd.DataFrame(rows)

# ========= STEP 7: YEAR RANGES BY CONTIGUOUS TRUE RANK =========
segments = []
current_rank = None
seg_start = None
seg_end = None

def max_dt(a, b):
    if pd.isna(a): return b
    if pd.isna(b): return a
    return max(a, b)

def min_dt(a, b):
    if pd.isna(a): return b
    if pd.isna(b): return a
    return min(a, b)

for _, row in df.iterrows():
    tr = row.get('true_rank')
    ds = row.get('date_start')
    de = row.get('date_end')
    if current_rank is None:
        current_rank = tr
        seg_start = ds
        seg_end = de
        continue
    if tr != current_rank:
        segments.append({'true_rank': current_rank, 'start': seg_start, 'end': seg_end})
        current_rank = tr
        seg_start = ds
        seg_end = de
    else:
        seg_start = min_dt(seg_start, ds)
        seg_end = max_dt(seg_end, de)

if current_rank is not None:
    segments.append({'true_rank': current_rank, 'start': seg_start, 'end': seg_end})

def fmt_year_range(start_dt, end_dt):
    if pd.isna(start_dt) and pd.isna(end_dt):
        return "Unknown"
    if pd.isna(start_dt):
        return f"–{end_dt.year}"
    sy = start_dt.year
    if pd.isna(end_dt):
        return f"{sy}–Present"
    ey = end_dt.year
    if ey < sy:
        return f"{sy}"
    return f"{sy}–{ey}"

year_ranges = [
    {'true_rank': seg['true_rank'], 'start': seg.get('start'), 'end': seg.get('end'),
     'year_range': fmt_year_range(seg.get('start'), seg.get('end'))}
    for seg in segments if seg.get('true_rank') is not None
]

# ========= STEP 8: LEARNING LOOP (vocab load + expansion utilities) =========
vocab = load_vocab()
role_exp = vocab.get("role_expansions", {})
loc_alias = vocab.get("location_aliases", {})

TEMP_DES_PATTERN = re.compile(r'\((?:TEMP|TEMPORARY|DES|DESIGNATE|DESIGNATED)\)', flags=re.IGNORECASE)

def _is_blankish(s: str) -> bool:
    return s is None or (isinstance(s, float) and pd.isna(s)) or str(s).strip().lower() in {"", "nan", "none", "null"}

def clean_role_token(txt: str) -> str:
    if _is_blankish(txt):
        return ""
    s = str(txt).strip()
    s = TEMP_DES_PATTERN.sub("", s)
    s = re.sub(r'\s+', ' ', s).strip()
    if s.lower() in {"nan", "none", "null", "-"}:
        return ""
    return s

def clean_loc_label(txt: str) -> str:
    if _is_blankish(txt):
        return ""
    s = str(txt).strip()
    s = re.sub(r'\s+', ' ', s).strip()
    if s.lower() in {"nan", "none", "null"}:
        return ""
    return s

unknown_role_tokens = set()
unknown_location_labels = set()

# === Smart casing for roles (keep common acronyms) ===
ROLE_ACRONYMS = {
    "PTU", "EU", "RCCC", "CCB", "CAPO", "HQCCC", "PCRO", "PPRB",
    "SDS", "DVIT", "DSDS", "PSU", "TFSU", "ADC", "DDC", "DC", "RI", "ES"
}
LOWERCASE_CONNECTORS = {"and", "of", "the", "in", "on", "for", "to", "with", "at"}

def smart_title_case_role(text: str) -> str:
    if not text: return text
    def fix_word(w: str, is_first: bool) -> str:
        ww = re.sub(r'[()\-/,]', '', w)
        if ww.upper() in ROLE_ACRONYMS:
            return w.upper()
        if not is_first and ww.lower() in LOWERCASE_CONNECTORS:
            return w.lower()
        # Title-case, but keep inner acronym segments
        return w[:1].upper() + w[1:].lower()
    parts = re.split(r'(\s+)', text)  # keep spaces
    out = []
    word_idx = 0
    for p in parts:
        if p.isspace():
            out.append(p)
        else:
            out.append(fix_word(p, is_first=(word_idx == 0)))
            word_idx += 1
    return ''.join(out)

# Canonicalization preferences (dedup synonyms -> full words)
CANON_SYNONYMS = {
    # OPS → Operations (and numbered)
    r'\bOPS\b': 'Operations',
    r'\bOPS\s*\((\d+)\)': r'Operations (\1)',
    r'^OPERATIONS\s*\((\d+)\)$': r'Operations (\1)',
    r'^OPERATIONS$': 'Operations',

    # HQCCC / HQCCC(OPS RM)
    r'\bHQCCC\s*\(OPS\s*RM\)\b': 'Headquarters Command and Control Centre (Operations Room)',
    r'\bHQCCC\b': 'Headquarters Command and Control Centre',
    r'^HEADQUARTERS COMMAND AND CONTROL CENTRE\s*\(OPERATIONS ROOM\)$': 'Headquarters Command and Control Centre (Operations Room)',
    r'^HEADQUARTERS COMMAND AND CONTROL CENTRE$': 'Headquarters Command and Control Centre',

    # Community relations
    r'\bCMU\s*REL\b': 'Community Relations',
    r'^COMMUNITY RELATIONS$': 'Community Relations',
    r'\bPCRO\b': 'Police Community Relations Office',
    r'^POLICE COMMUNITY RELATIONS OFFICE$': 'Police Community Relations Office',

    # PSU / TFSU / DVIT / DIVT / SDS / DSDS
    r'\bPSU\s*(\d+)\b': r'Patrol Sub-unit \1',
    r'^PATROL\s+SUB-UNIT\s*(\d+)$': r'Patrol Sub-unit \1',
    r'\bTFSU\b': 'Task Force Sub-unit',
    r'^TASK\s+FORCE\s+SUB-UNIT$': 'Task Force Sub-unit',
    r'\bDVIT\s*(\d+)\b': r'Divisional Investigation Team \1',
    r'\bDIVT\s*(\d+)\b': r'Divisional Investigation Team \1',
    r'^DIVISIONAL\s+INVESTIGATION\s+TEAM\s*(\d+)$': r'Divisional Investigation Team \1',
    r'\bSDS\s*(\d+)\b': r'Special Duties Squad \1',
    r'\bDSDS\s*(\d+)\b': r'District Special Duties Squad \1',
    r'^SPECIAL DUTIES SQUAD\s*(\d+)$': r'Special Duties Squad \1',

    # Command/Commander short-hands (Platoon consolidation)
    r'\bPLN\s*(\d+)\s*(?:CDR|COMMANDER)\b': r'Platoon \1 Commander',
    r'\bPLATOON\s*(\d+)\s*(?:CDR|COMMANDER)\b': r'Platoon \1 Commander',
    r'\bPLN\s*(\d+)\b': r'Platoon \1',
    r'^PLATOON\s*(\d+)$': r'Platoon \1',
    # Standalone CDR is too vague; drop it to avoid noise duplicates
    r'^(?:CDR)$': '',

    # Role words to consistent case
    r'^ADMINISTRATION$': 'Administration',
    r'^GENERAL$': 'General',
    r'^FIELD$': 'Field',
    r'^SECURITY ADVISORY SECTION$': 'Security Advisory Section',
    r'^ARCHITECTURAL LIAISON$': 'Architectural Liaison',
    r'^PUBLICITY$': 'Publicity',
    r'^ACH\s*LIA$': 'Architectural Liaison',
    r'^PUB$': 'Publicity',
    r'^SERGEANT FORCE PROMOTION ASSESSMENT TEAM$': 'Sergeant Force Promotion Assessment Team',
    r'^SECURITY COMPANY AND GUARDING SERVICES BILL-POLICE INSPECTION TEAM$': 'Security Company and Guarding Services Bill-Police Inspection Team',
    r'^SDVC$': 'Sub-divisional Commander',

    # CRM -> Crime variants (to avoid duplicates)
    r'\bCRM\s*\((\d+)\)\b': r'Crime (\1)',
    r'\bCRM\b': 'Crime',
}

def canonicalize_role(name: str) -> str:
    s = clean_role_token(name)
    if not s:
        return ""
    # Run synonym expansions
    for pattern, repl in CANON_SYNONYMS.items():
        s = re.sub(pattern, repl, s, flags=re.IGNORECASE)
    s = re.sub(r'\s+', ' ', s).strip()
    # Smart casing (keep acronyms), but do not uppercase everything
    s = smart_title_case_role(s)
    return s

def expand_role(token: str) -> str:
    t = clean_role_token(token)
    if not t:
        return ""
    if t in role_exp:
        return canonicalize_role(role_exp[t])
    t_up = t.upper()
    if t_up in role_exp:
        return canonicalize_role(role_exp[t_up])

    # Inline patterns
    m = re.search(r'\bD(V)?IVT\s*(\d+)\b', t_up)
    if m:
        return canonicalize_role(f"Divisional Investigation Team {m.group(2)}")

    m = re.search(r'\bPSU\s*(\d+)\b', t_up)
    if m:
        return canonicalize_role(f"Patrol Sub-unit {m.group(1)}")

    m = re.search(r'\bDS?DS\s*(\d+)\b', t_up)
    if m:
        if t_up.startswith('DSDS'):
            return canonicalize_role(f"District Special Duties Squad {m.group(1)}")
        else:
            return canonicalize_role(f"Special Duties Squad {m.group(1)}")

    if 'HQCCC' in t_up:
        if 'OPS' in t_up and 'RM' in t_up:
            return canonicalize_role('Headquarters Command and Control Centre (Operations Room)')
        return canonicalize_role('Headquarters Command and Control Centre')

    s = canonicalize_role(t_up)
    if s.upper() == t_up:
        unknown_role_tokens.add(t)
    return s

def normalize_location(label: str) -> str:
    l = clean_loc_label(label)
    if not l:
        return ""
    if l.upper() == "LEAVE RESERVE":
        return ""
    if l in loc_alias:
        return loc_alias[l]
    l_up = l.upper()
    if l_up in loc_alias:
        return loc_alias[l_up]
    unknown_location_labels.add(l)
    return l

def is_rank_text(text: str) -> bool:
    if not text:
        return False
    mapped = map_rank(text)
    return mapped in rank_index

# ========= STEP 9: BUILD ENHANCED RANGES (locations + roles per location) =========
def consolidate_row_roles(roles_out: list[str]) -> list[str]:
    """
    Collapse near-duplicates with preference for fuller wording.
    - If both 'Platoon N' and 'Platoon N Commander' exist, keep only 'Platoon N Commander'.
    """
    # Index roles by normalized key
    keep = set(roles_out)
    # Map Platoon patterns
    platoon_nums = {}
    for r in roles_out:
        m_cmd = re.match(r'^Platoon\s+(\d+)\s+Commander$', r, flags=re.IGNORECASE)
        m_simple = re.match(r'^Platoon\s+(\d+)$', r, flags=re.IGNORECASE)
        if m_cmd:
            platoon_nums.setdefault(int(m_cmd.group(1)), {"cmd": None, "plain": None})
            platoon_nums[int(m_cmd.group(1))]["cmd"] = r
        if m_simple:
            platoon_nums.setdefault(int(m_simple.group(1)), {"cmd": None, "plain": None})
            platoon_nums[int(m_simple.group(1))]["plain"] = r

    for n, d in platoon_nums.items():
        if d.get("cmd") and d.get("plain"):
            keep.discard(d["plain"])

    # Remove empty strings if ever present
    keep.discard("")
    return list(keep)

enhanced_ranges = []
for seg in year_ranges:
    tr = seg['true_rank']
    start_dt = seg['start']
    end_dt = seg['end']

    mask = (df['true_rank'] == tr)
    if pd.notna(start_dt):
        mask &= (df['date_start'].isna() | (df['date_start'] >= start_dt))
    if pd.notna(end_dt):
        mask &= (df['date_end'].isna() | (df['date_end'] <= end_dt))

    sub = df.loc[mask].copy()

    # Preferred location: Location (Description) -> fallback Location -> normalize/alias
    loc_desc = sub['location_desc']
    loc_series = loc_desc.where(
        loc_desc.notna() & (loc_desc.astype(str).str.strip() != ''),
        sub['location']
    ).apply(normalize_location)

    # Extract roles per row
    def extract_roles_from_row(r):
        roles_out = []

        # 1) Prefer Designation Description
        dd_raw = r.get('designation_desc') or ''
        dd = canonicalize_role(expand_role(dd_raw))
        if dd:
            roles_out.append(dd)

        # 2) Add Designation only if description absent/different and not rank
        d_raw = r.get('designation') or ''
        if d_raw and not is_rank_text(d_raw):
            d = canonicalize_role(expand_role(d_raw))
            if d and (not dd or d.casefold() != dd.casefold()):
                roles_out.append(d)

        # 3) Post Type only if non-rank
        pt = r.get('post_type') or ''
        if pt and not is_rank_text(pt):
            p = canonicalize_role(expand_role(pt))
            if p:
                roles_out.append(p)

        # Row-level dedup (case-insensitive)
        seen = set()
        clean_list = []
        for x in roles_out:
            if not x or x.lower() in {"nan", "none", "null"}:
                continue
            key = x.casefold()
            if key not in seen:
                seen.add(key)
                clean_list.append(x)

        # Row-level consolidation (e.g., Platoon N vs Platoon N Commander)
        clean_list = consolidate_row_roles(clean_list)
        return clean_list

    sub['role_list'] = sub.apply(extract_roles_from_row, axis=1)

    # Group roles by location (case-insensitive dedup within each location)
    roles_by_loc = {}
    seen_by_loc = {}

    for l, roles_here in zip(loc_series, sub['role_list']):
        if not l:
            continue
        roles_by_loc.setdefault(l, [])
        seen_by_loc.setdefault(l, set())
        for role in roles_here:
            if not role or role.upper() == "LEAVE RESERVE":
                continue
            role = canonicalize_role(role)  # ensure same pipeline
            key = role.casefold()
            if key not in seen_by_loc[l]:
                roles_by_loc[l].append(role)
                seen_by_loc[l].add(key)

    # ===== Enhanced Division→District merge (DIV, DIVISION -> DIST, DISTRICT) =====
    def loc_base_and_type(label: str):
        """
        Returns (base_upper, type) where type in {'DISTRICT','DIVISION','SUB-DIVISION',None}
        """
        s = str(label or "")
        s_up = re.sub(r'\s+', ' ', s).strip().upper()

        # Identify type
        typ = None
        if re.search(r'\bSUB[-\s]?DIVISION\b', s_up):
            typ = 'SUB-DIVISION'
        elif re.search(r'\bDIV(?:ISION)?\b', s_up):
            typ = 'DIVISION'
        elif re.search(r'\bDIST(?:RICT)?\b', s_up):
            typ = 'DISTRICT'

        # Remove the tokens to get base
        base = s_up
        base = re.sub(r'\bSUB[-\s]?DIVISION\b', '', base)
        base = re.sub(r'\bDIV(?:ISION)?\b', '', base)
        base = re.sub(r'\bDIST(?:RICT)?\b', '', base)
        base = re.sub(r'\s+', ' ', base).strip()
        return base, typ

    # Build mapping of DIVISION -> DISTRICT for same base
    all_loc_names = set(roles_by_loc.keys())
    base_to_types = {}
    for loc_name in list(all_loc_names):
        base, typ = loc_base_and_type(loc_name)
        base_to_types.setdefault(base, set()).add(typ)

    division_to_district = {}
    for loc_name in list(all_loc_names):
        base, typ = loc_base_and_type(loc_name)
        if typ == 'DIVISION' and 'DISTRICT' in base_to_types.get(base, set()):
            # which concrete district name?
            # Find a matching district label with same base
            for other in all_loc_names:
                b2, t2 = loc_base_and_type(other)
                if b2 == base and t2 == 'DISTRICT':
                    division_to_district[loc_name] = other
                    break

    # Merge roles
    merged_roles_by_loc = {}
    merged_seen = {}

    for loc_name, roles in roles_by_loc.items():
        target = division_to_district.get(loc_name, loc_name)
        merged_roles_by_loc.setdefault(target, [])
        merged_seen.setdefault(target, set())
        for rname in roles:
            key = rname.casefold()
            if key not in merged_seen[target]:
                merged_roles_by_loc[target].append(rname)
                merged_seen[target].add(key)

    # Deduplicate locations preserving order, then rebuild after merge
    unique_locs = []
    seen_locs = set()
    for l in loc_series.tolist():
        if not l:
            continue
        if l not in seen_locs:
            seen_locs.add(l)
            unique_locs.append(l)

    new_unique_locs = []
    seen_after = set()
    for loc_name in unique_locs:
        target = division_to_district.get(loc_name, loc_name)
        if target not in seen_after:
            seen_after.add(target)
            new_unique_locs.append(target)

    enhanced_ranges.append({
        'true_rank': tr,
        'year_range': seg['year_range'],
        'locations': new_unique_locs,
        'roles_by_location': merged_roles_by_loc
    })

# ========= STEP 10: data0, data1, ... convenience =========
data_arrays = []
for _, row in df.iterrows():
    data_arrays.append(row.tolist())
for i, arr in enumerate(data_arrays):
    globals()[f"data{i}"] = arr

# ========= STEP 11: OUTPUTS =========
print("\n=== ALL rows with computed ranks ===")
show_cols = [c for c in [
    'date_start', 'date_end',
    'post_type', 'post_type_desc',
    'designation', 'designation_desc',
    'reported_rank', 'acting_flag', 'true_rank',
    'location', 'location_desc'
] if c in df.columns]
if show_cols:
    print(df[show_cols].to_string(index=False))
else:
    print("(No displayable columns found)")

print("\n=== True Rank Year Ranges (contiguous) + Locations & Roles ===")
for item in enhanced_ranges:
    print(f"{item['true_rank']}: {item['year_range']}")
    if not item['locations']:
        continue
    for loc in item['locations']:
        print(f"  {loc}")
        roles = item['roles_by_location'].get(loc, [])
        if roles:
            for rname in roles:
                print(f"    - {rname}")

print("\nTotal rows of data:", len(data_arrays))
if len(data_arrays) > 0:
    print("Example: data0 =", data0)
if len(data_arrays) > 1:
    print("Example: data1 =", data1)

# ========= STEP 12: Unknowns export + save vocab =========
if EXPORT_UNKNOWNS:
    unknown_payload = {
        "role_tokens": sorted(t for t in unknown_role_tokens if t),
        "location_labels": sorted(l for l in unknown_location_labels if l)
    }
    with open(UNKNOWN_PATH, "w", encoding="utf-8") as f:
        json.dump(unknown_payload, f, ensure_ascii=False, indent=2)
    print(f"\n[Info] Exported unknown tokens to {UNKNOWN_PATH.resolve()}")

save_vocab(vocab)
# ========= STEP 13: GENERATE WORD DOCUMENT =========
def generate_word_document(enhanced_ranges, output_filename="HKPF_Posting_Summary.docx"):
    """
    Generate a Word document with a table format.
    Table columns: Year Range | Rank | Posting Location & Roles
    """
    doc = Document()
    
    # Rank expansion mapping (without "of Police")
    rank_full_names = {
        'PC': 'Police Constable',
        'SPC': 'Senior Police Constable',
        'SGT': 'Sergeant',
        'SSGT': 'Station Sergeant',
        'PI': 'Probationary Inspector',
        'IP': 'Inspector',
        'SIP': 'Senior Inspector',
        'IP/SIP': 'Inspector / Senior Inspector',
        'CIP': 'Chief Inspector',
        'SP': 'Superintendent',
        'SSP': 'Senior Superintendent',
        'CSP': 'Chief Superintendent',
        'ACP': 'Assistant Commissioner',
        'SACP': 'Senior Assistant Commissioner',
        'DCP': 'Deputy Commissioner',
        'CP': 'Commissioner',
    }
    
    # Helper function to shade cell
    def shade_cell(cell, color):
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Create table with 3 columns
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Set column widths (narrow for Year Range and Rank, wide for Posting Info)
    table.columns[0].width = Inches(1.0)  # Year Range - narrow
    table.columns[1].width = Inches(1.2)  # Rank - narrow
    table.columns[2].width = Inches(4.3)  # Posting Location & Roles - wide
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Year Range"
    header_cells[1].text = "Rank"
    header_cells[2].text = "Posting Location & Roles"
    
    # Format header row - bold, size 13, light blue background
    for cell in header_cells:
        shade_cell(cell, "ADD8E6")  # Light blue
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(13)
    
    # Add data rows
    for item in enhanced_ranges:
        year_range = item['year_range']
        rank = item['true_rank']
        locations = item['locations']
        roles_by_location = item['roles_by_location']
        
        # Convert rank acronym to full name
        rank_display = rank_full_names.get(rank, rank)
        
        # Build posting info text
        posting_info = []
        if locations:
            for loc in locations:
                posting_info.append(f"{loc}")
                roles = roles_by_location.get(loc, [])
                if roles:
                    for role in roles:
                        posting_info.append(f"  • {role}")
        else:
            posting_info.append("(No locations recorded)")
        
        posting_text = "\n".join(posting_info)
        
        # Add row
        row_cells = table.add_row().cells
        row_cells[0].text = year_range
        row_cells[1].text = rank_display
        row_cells[2].text = posting_text
        
        # Format data rows - white background, size 13
        for cell in row_cells:
            shade_cell(cell, "FFFFFF")  # White
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(13)
    
    # Save document
    doc.save(output_filename)
    return output_filename

# Generate the Word document
docx_file = generate_word_document(enhanced_ranges)
print(f"\n[Success] Word document generated: {docx_file}")
print(f"Location: {Path(docx_file).resolve()}")