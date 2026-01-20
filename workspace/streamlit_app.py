import streamlit as st
import pandas as pd
import re
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
import io

st.set_page_config(page_title="HKPF Posting Summary Generator", layout="wide")

st.title("📊 HKPF Posting Summary Generator")
st.write("Upload an Excel file to generate a professional posting summary Word document")

# ========= VOCAB + CONFIGURATION =========
STARTER_ROLE_EXPANSIONS = {
    "CSP": "Chief Superintendent",
    "SSP": "Senior Superintendent",
    "SP": "Superintendent",
    "CIP": "Chief Inspector",
    "SIP": "Senior Inspector",
    "IP": "Inspector",
    "PI": "Probationary Inspector",
    "SSGT": "Station Sergeant",
    "SGT": "Sergeant",
    "SPC": "Senior Police Constable",
    "PC": "Police Constable",
    "DC": "District Commander",
    "DDC": "Deputy District Commander",
    "ADC": "Assistant District Commander",
    "ADM": "Administration",
    "A&S": "Administration and Support",
    "CRM": "Crime",
    "ES": "Efficiency Studies",
    "RI": "Research and Inspections",
    "CTRL": "Command and Control (Control Room)",
    "GEN": "General",
    "FLD": "Field",
    "PSU 1": "Patrol Sub-unit 1",
    "PSU 2": "Patrol Sub-unit 2",
    "PSU 3": "Patrol Sub-unit 3",
    "PSU 4": "Patrol Sub-unit 4",
    "TFSU": "Task Force Sub-unit",
    "DVIT 1": "Divisional Investigation Team 1",
    "DVIT 2": "Divisional Investigation Team 2",
    "DVIT 3": "Divisional Investigation Team 3",
    "DVIT 4": "Divisional Investigation Team 4",
    "DVIT 5": "Divisional Investigation Team 5",
    "DVIT 6": "Divisional Investigation Team 6",
    "DVIT 7": "Divisional Investigation Team 7",
    "DVIT 8": "Divisional Investigation Team 8",
    "SDS 1": "Special Duties Squad 1",
    "DSDS 2": "District Special Duties Squad 2",
    "SYMPOSIUM": "Symposium",
    "RPC TRG (INTAKE)": "Recruit Police Constable Training (Intake)",
    "CS&INT": "Counterfeit, Support and Intelligence",
    "INP 2": "Inspection 2",
    "AUX": "Auxiliary",
    "SCIU": "Security Company and Guarding Services Bill-Police Inspection Team",
    "SA": "Security Advisory Section",
    "PCRO": "Police Community Relations Office",
}

STARTER_LOCATION_ALIASES = {
    # Districts (normalized forms) - VERIFIED
    "CDIST": "CENTRAL DISTRICT",
    "CDIV": "CENTRAL DIVISION",
    "WDIST": "WESTERN DISTRICT",
    "WDIV": "WESTERN DIVISION",
    "EDIST": "EASTERN DISTRICT",
    "MKDIST": "MONG KOK DISTRICT",
    "NPDIST": "NORTH POINT DISTRICT",
    "NPDIV": "NORTH POINT DIVISION",
    "SMPDIST": "SAU MAU PING DISTRICT",
    "TWDIST": "TSUEN WAN DISTRICT",
    "TPDIST": "TUEN MUN DISTRICT",
    "TPDIV": "TUEN MUN DIVISION",
    "WTDIST": "WAN CHAI DISTRICT",
    "WTDIV": "WAN CHAI DIVISION",
    "STDIST": "SHA TIN DISTRICT",
    "STDIV": "SHA TIN DIVISION",
    "YLDIST": "YUEN LONG DISTRICT",
    "YLDIV": "YUEN LONG DIVISION",
    "TPDIST2": "TAI PO DISTRICT",
    "TPDIV2": "TAI PO DIVISION",
    "KQDIST": "KWAI TSING DISTRICT",
    "KQDIV": "KWAI TSING DIVISION",
    "STDIV3": "STANLEY DIVISION",
    
    # Divisions
    "WCH DIV": "WAN CHAI DIVISION",
    "WCH DIST": "WAN CHAI DISTRICT",
    "STYSDIV": "STANLEY DIVISION",
    
    # Regional Codes (VERIFIED)
    "HKI": "HONG KONG ISLAND",
    "KW": "KOWLOON WEST",
    "KE": "KOWLOON EAST",
    "NTN": "NEW TERRITORIES NORTH",
    "NTS": "NEW TERRITORIES SOUTH",
    
    # Regional Command and Control (VERIFIED)
    "RCCC HKI": "REGIONAL COMMAND AND CONTROL CENTRE HONG KONG ISLAND",
    "RCCC KW": "REGIONAL COMMAND AND CONTROL CENTRE KOWLOON WEST",
    "RCCC KE": "REGIONAL COMMAND AND CONTROL CENTRE KOWLOON EAST",
    "RCCC NTN": "REGIONAL COMMAND AND CONTROL CENTRE NEW TERRITORIES NORTH",
    "RCCC NTS": "REGIONAL COMMAND AND CONTROL CENTRE NEW TERRITORIES SOUTH",
    
    # Emergency and Tactical Units (VERIFIED)
    "EU HKI": "EMERGENCY UNIT HONG KONG ISLAND",
    "EU KW": "EMERGENCY UNIT KOWLOON WEST",
    "EU KE": "EMERGENCY UNIT KOWLOON EAST",
    "EU NTN": "EMERGENCY UNIT NEW TERRITORIES NORTH",
    "EU NTS": "EMERGENCY UNIT NEW TERRITORIES SOUTH",
    "PTU": "POLICE TACTICAL UNIT",
    "PTU A": "POLICE TACTICAL UNIT (A COMPANY)",
    "PTU W": "POLICE TACTICAL UNIT (WEST COMPANY)",
    
    # Traffic Units
    "T HKI": "TRAFFIC HONG KONG ISLAND",
    "T KW": "TRAFFIC KOWLOON WEST",
    "T KE": "TRAFFIC KOWLOON EAST",
    "T NTN": "TRAFFIC NEW TERRITORIES NORTH",
    "T NTS": "TRAFFIC NEW TERRITORIES SOUTH",
    
    # Crime Units by Region
    "CRM HKI": "CRIME BUREAU HONG KONG ISLAND",
    "CRM KW": "CRIME BUREAU KOWLOON WEST",
    "CRM KE": "CRIME BUREAU KOWLOON EAST",
    "CRM NTN": "CRIME BUREAU NEW TERRITORIES NORTH",
    "CRM NTS": "CRIME BUREAU NEW TERRITORIES SOUTH",
    
    # Operations Units by Region
    "OPS HKI": "OPERATIONS HONG KONG ISLAND",
    "OPS KW": "OPERATIONS KOWLOON WEST",
    "OPS KE": "OPERATIONS KOWLOON EAST",
    "OPS NTN": "OPERATIONS NEW TERRITORIES NORTH",
    "OPS NTS": "OPERATIONS NEW TERRITORIES SOUTH",
    
    # Regional Headquarters
    "KW RHQ": "KOWLOON WEST REGIONAL HEADQUARTERS",
    "KE RHQ": "KOWLOON EAST REGIONAL HEADQUARTERS",
    "HKI RHQ": "HONG KONG ISLAND REGIONAL HEADQUARTERS",
    
    # Bureaus and Branches (VERIFIED)
    "CCB": "COMMERCIAL CRIME BUREAU",
    "C DIV CCB": "COMMERCIAL CRIME BUREAU",
    "C DIVISION COMMERCIAL CRIME BUREAU": "COMMERCIAL CRIME BUREAU",
    "CPB": "CRIME PREVENTION BUREAU",
    "PPRB": "POLICE PUBLIC RELATIONS BRANCH",
    
    # Police Offices (VERIFIED)
    "CAPO": "COMPLAINTS AGAINST POLICE OFFICE",
    "CAPO HKI": "COMPLAINTS AGAINST POLICE OFFICE HONG KONG ISLAND",
    
    # Training and Support (VERIFIED)
    "PTS": "POLICE TRAINING SCHOOL",
    "PC TRG": "POLICE CONSTABLE TRAINING DIVISION",
    "IST": "IN-SERVICE TRAINING",
    "TRVE SUP": "TRAINING RESERVE SUPPORT WING",
    
    # Wings and Headquarters (VERIFIED)
    "SQ": "SERVICE QUALITY WING",
    "SUPPORT": "SUPPORT WING",
    "OPS": "OPERATIONS WING",
    "HQCCC": "HEADQUARTERS COMMAND AND CONTROL CENTRE",
    
    # Intelligence and Anti-Triad Units (VERIFIED)
    "RATU KW": "REGIONAL ANTI TRIAD UNIT KOWLOON WEST",
    "RATU KE": "REGIONAL ANTI TRIAD UNIT KOWLOON EAST",
    "RATU NTN": "REGIONAL ANTI TRIAD UNIT NEW TERRITORIES NORTH",
    "RIU KW": "REGIONAL INTELLIGENCE UNIT KOWLOON WEST",
    "RIU KE": "REGIONAL INTELLIGENCE UNIT KOWLOON EAST",
    "RIU NTN": "REGIONAL INTELLIGENCE UNIT NEW TERRITORIES NORTH",
}

# ========= CORE PROCESSING FUNCTIONS =========
rank_order = ['PC', 'SPC', 'SGT', 'SSGT', 'PI', 'IP', 'SIP', 'CIP', 'SP', 'SSP', 'CSP', 'ACP', 'SACP', 'DCP', 'CP']
rank_index = {r: i for i, r in enumerate(rank_order)}
rank_index['IP/SIP'] = rank_index['IP']

rank_map = {
    'pc': 'PC', 'police constable': 'PC',
    'spc': 'SPC', 'senior police constable': 'SPC',
    'sgt': 'SGT', 'sergeant': 'SGT',
    'ssgt': 'SSGT', 'station sergeant': 'SSGT',
    'pi': 'PI', 'probationary inspector': 'PI', 'probationary inspector of police': 'PI',
    'ip': 'IP', 'insp': 'IP', 'inspector': 'IP', 'inspector of police': 'IP',
    'sip': 'SIP', 'sr insp': 'SIP', 'sen insp': 'SIP', 'senior insp': 'SIP', 'senior inspector': 'SIP', 'senior inspector of police': 'SIP',
    'cip': 'CIP', 'ch insp': 'CIP', 'chief insp': 'CIP', 'chief inspector': 'CIP', 'chief inspector of police': 'CIP',
    'sp': 'SP', 'superintendent': 'SP', 'superintendent of police': 'SP',
    'ssp': 'SSP', 'senior superintendent': 'SSP', 'senior superintendent of police': 'SSP',
    'csp': 'CSP', 'chief superintendent': 'CSP', 'chief superintendent of police': 'CSP',
}

acting_tokens_pattern = re.compile(r'\b(acting|actg|a/|ag\.|temp|temporary|acting up)\b', flags=re.IGNORECASE)

def looks_like_ip_sip(text: str) -> bool:
    s = str(text or "").lower()
    s = acting_tokens_pattern.sub('', s)
    s = s.replace('\\', '/')
    s = re.sub(r'[().,;]', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    s2 = s
    s2 = re.sub(r'\bsenior\s+inspector(?:\s+of\s+police)?\b', 'sip', s2)
    s2 = re.sub(r'\bsr\s+insp\b', 'sip', s2)
    s2 = re.sub(r'\bsen\s+insp\b', 'sip', s2)
    s2 = re.sub(r'\binspector(?:\s+of\s+police)?\b', 'ip', s2)
    s2 = re.sub(r'\binsp\b', 'ip', s2)
    tokens = set(re.split(r'[^a-z/]+', s2))
    tokens.discard('')
    return ('ip' in tokens and 'sip' in tokens) or ('ip/sip' in s2)

def map_rank(text: str):
    if looks_like_ip_sip(text):
        return 'IP/SIP'
    s = str(text or "").strip().lower()
    s = acting_tokens_pattern.sub('', s)
    s = s.replace('\\', '/')
    s = re.sub(r'[().,;]', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    if not s:
        return None
    if s in rank_map:
        v = rank_map[s]
    else:
        v = None
        for k, val in rank_map.items():
            if re.search(r'\b' + re.escape(k) + r'\b', s):
                v = val
                break
    if v in {'IP', 'SIP'}:
        return 'IP/SIP'
    return v

def is_acting(row) -> bool:
    fields = [
        str(row.get('designation', '') or ''),
        str(row.get('designation_desc', '') or ''),
        str(row.get('post_type', '') or ''),
        str(row.get('post_type_desc', '') or ''),
    ]
    haystack = ' || '.join(fields)
    return bool(acting_tokens_pattern.search(haystack))

def detect_header_row(df_raw: pd.DataFrame) -> int:
    target_headers = {'date start', 'date end', 'post type'}
    header_row = None
    for i in range(min(60, len(df_raw))):
        row_vals = df_raw.iloc[i].astype(str).str.strip().str.lower().tolist()
        if target_headers.issubset(set(row_vals)):
            header_row = i
            break
    if header_row is None:
        raise ValueError("Couldn't find header row with Date Start/Date End/Post Type.")
    return header_row

def snake(s: str) -> str:
    return re.sub(r'\s+', '_', str(s).strip().lower())

def _is_blankish(s: str) -> bool:
    return s is None or (isinstance(s, float) and pd.isna(s)) or str(s).strip().lower() in {"", "nan", "none", "null"}

def clean_role_token(txt: str) -> str:
    if _is_blankish(txt):
        return ""
    s = str(txt).strip()
    s = re.sub(r'\s+', ' ', s).strip()
    if s.lower() in {"nan", "none", "null", "-"}:
        return ""
    return s

def clean_loc_label(txt: str) -> str:
    if _is_blankish(txt):
        return ""
    s = str(txt).strip()
    s = re.sub(r'\s+', ' ', s).strip()
    if s.lower() in {"nan", "none", "null"}:
        return ""
    return s

def normalize_location(label: str, loc_alias) -> str:
    l = clean_loc_label(label)
    if not l:
        return ""
    if l.upper() == "LEAVE RESERVE":
        return ""
    if l in loc_alias:
        return loc_alias[l]
    l_up = l.upper()
    if l_up in loc_alias:
        return loc_alias[l_up]
    return l

def is_rank_text(text: str) -> bool:
    if not text:
        return False
    mapped = map_rank(text)
    return mapped in rank_index

# Enhanced canonicalization with comprehensive synonym rules
CANON_SYNONYMS = {
    # RI / Research and Inspections (multiple variants)
    r'^RI$': 'Research and Inspections',
    r'\bRI\b': 'Research and Inspections',
    r'\bR\s*&\s*I\b': 'Research and Inspections',
    r'\bR\s*and\s+I\b': 'Research and Inspections',
    r'\bResearch\s+&\s+Inspections\b': 'Research and Inspections',
    
    # T / Traffic (unit/department acronym - only standalone T)
    r'^T$': 'Traffic',
    
    # Inspection / INP / Inspection variants
    r'^INP(?:\s+2)?$': 'Inspection',
    r'\bINP\b': 'Inspection',
    r'\bInspection\s*\d+': 'Inspection',
    
    # CRM / Crime
    r'^CRM$': 'Crime',
    r'\bCRM\s*\((\d+)\)': r'Crime',
    r'\bCRM\b': 'Crime',
    
    # CS&INT / Counterfeit Support and Intelligence (multiple variants)
    r'^CS\s*&\s*INT$': 'Counterfeit, Support and Intelligence',
    r'\bCS\s*&\s*INT\b': 'Counterfeit, Support and Intelligence',
    r'\bCounterfeit\s*,?\s*Support\s+and\s+Intelligence\b': 'Counterfeit, Support and Intelligence',
    
    # RPC TRG / Recruit Police Constable Training
    r'\bRPC\s+TRG\s*\(INTAKE\)': 'Recruit Police Constable Training',
    r'\bRPC\s+TRG\b': 'Recruit Police Constable Training',
    
    # Administration variants
    r'^ADM$': 'Administration',
    r'\bADM\b': 'Administration',
    
    # Control Room
    r'^CTRL$': 'Command and Control (Control Room)',
    r'\bCTRL\b': 'Command and Control (Control Room)',
    
    # A&S / Administration and Support
    r'\bA\s*&\s*S\b': 'Administration and Support',
    
    # OSSUC / Operations Sub-unit Commander
    r'\bOSSUC\b': 'Operations Sub-unit Commander',
    
    # ADVC / Assistant Divisional Commander
    r'\bADVC\b': 'Assistant Divisional Commander',
    
    # ASSUC / Administration Sub-unit Commander
    r'\bASSUC\b': 'Administration Sub-unit Commander',
    
    # MESUC / Miscellaneous Enquiries Sub-unit Commander
    r'\bMESUC\b': 'Miscellaneous Enquiries Sub-unit Commander',
    
    # MESU / Miscellaneous Enquiries Sub-unit (including PSUC variant)
    r'\bMESU\b': 'Miscellaneous Enquiries Sub-unit',
    r'\bPSUC\b': 'Patrol Sub-unit Commander',
    
    # OPS variants
    r'\bOPS\s*\((\d+)\)': r'Operations',
    r'\bOPS\b': 'Operations',
    
    # HQCCC variants (unify all to one canonical form)
    r'\bHQCCC\s*\((?:OPS\s*RM|Ops\s*Rm)\)': 'Headquarters Command and Control Centre (Operations Room)',
    r'\bHeadquarters\s+Command\s+and\s+Control\s+Centre\s+(?:Operations|operations)(?:\s+)?(?:Rm|Room|RM)': 'Headquarters Command and Control Centre (Operations Room)',
    r'\bHQCCC\s*\(operations\s+Room\)': 'Headquarters Command and Control Centre (Operations Room)',
    r'\bHQCCC\b': 'Headquarters Command and Control Centre',
    
    # DVIT / DIVT variants
    r'\bDIVT\s*(\d+)': r'Divisional Investigation Team',
    r'\bDVIT\s*(\d+)': r'Divisional Investigation Team',
    r'\bDivisional\s+Investigation\s+Team\s*(\d+)': r'Divisional Investigation Team',
    
    # PSU variants
    r'\bPSU\s*(\d+)': r'Patrol Sub-unit',
    r'\bPatrol\s+Sub[-\s]?unit\s*(\d+)': r'Patrol Sub-unit',
    
    # SDS / DSDS variants
    r'\bDSDS\s*(\d+)': r'District Special Duties Squad',
    r'\bD?SDS\s*(\d+)': r'Special Duties Squad',
    
    # TFSU
    r'\bTFSU\b': 'Task Force Sub-unit',
    
    # PCRO / Police Community Relations Office
    # Community relations (force all variants to canonical form)
    r'\bCMU\s*REL\b': 'Community Relations',
    r'^CMU REL$': 'Community Relations',
    r'^C M U REL$': 'Community Relations',
    r'^COMMUNITY RELATIONS$': 'Community Relations',
    r'^COMM REL$': 'Community Relations',
    r'\bPCRO\b': 'Police Community Relations Office',
    
    # ES / Efficiency Studies
    r'^ES$': 'Efficiency Studies',
    r'\bES\b': 'Efficiency Studies',
    
    # GEN / General
    r'^GEN$': 'General',
    r'\bGEN\b': 'General',
    
    # FLD / Field
    r'^FLD$': 'Field',
    r'\bFLD\b': 'Field',
    
    # ADC / Assistant Divisional Commander
    r'^ADC$': 'Assistant Divisional Commander',
    r'\bADC\b': 'Assistant Divisional Commander',
    
    # DDC / Deputy District Commander
    r'^DDC$': 'Deputy District Commander',
    r'\bDDC\b': 'Deputy District Commander',
    
    # DC / District Commander (but not when part of another word)
    r'^DC$': 'District Commander',
    r'\bDC\b(?!\w)': 'District Commander',
    
    # Platoon/Commander variants (standardize, strip numbers)
    r'\bPLN\s*(\d+)\b': 'Platoon Commander',
    r'\bPLATOON\s*(\d+)\b': 'Platoon Commander',
    r'\bCDR\s+PLN\s*(\d+)\b': 'Platoon Commander',
    r'\bCDR\b': 'Commander',
    r'\bPlatoon\s*(\d+)\s+Commander\b': 'Platoon Commander',
    
    # Symposium
    r'^SYMPOSIUM$': 'Symposium',
    r'\bSYMPOSIUM\b': 'Symposium',
    
    # Team variants (strip numbers)
    r'\bTEAM\s*\d+[A-Z]?\b': 'Team',
    r'\bTeam\s*\d+[A-Z]?\b': 'Team',
    
    # Common abbreviations that appear in role data
    r'^ACH\s*LIA$': 'Architectural Liaison',
    r'\bACH\s*LIA\b': 'Architectural Liaison',
    r'^PUB$': 'Publicity',
    r'\bPUB\b': 'Publicity',
    r'^SA$': 'Security Advisory Section',
    r'\bSA\b': 'Security Advisory Section',
    
    # Strip trailing numbers from all role names (Squad 1 → Squad, Team 2 → Team, etc.)
    r'Special Duties Squad\s+\d+': 'Special Duties Squad',
    r'District Special Duties Squad\s+\d+': 'District Special Duties Squad',
    r'Divisional Investigation Team\s+\d+': 'Divisional Investigation Team',
    r'Patrol Sub-unit\s+\d+': 'Patrol Sub-unit',
    r'Platoon\s+\d+': 'Platoon',
    r'Operations\s+\d+': 'Operations',
    r'(\w+\s+)*(\w+)\s+\d+$': r'\2',  # Catch-all: remove trailing numbers from any role
}

ROLE_ACRONYMS = {
    "HQCCC", "PTU", "EU", "RCCC", "CCB", "CAPO", "PCRO", "PPRB", 
    "DVIT", "PSU", "SDS", "DSDS", "RIU", "RATU", "ADC", "DDC", "DC", "RI", "ES", "OPS", "CS"
}

def smart_title_case_role(text: str) -> str:
    """Title case while preserving known acronyms (all caps)."""
    if not text:
        return text
    def fix_word(w: str, is_first: bool) -> str:
        ww = re.sub(r'[()\-/,]', '', w)
        if ww.upper() in ROLE_ACRONYMS:
            return w.upper()
        if not is_first and ww.lower() in {"and", "of", "the", "in", "on", "for", "to", "with", "at"}:
            return w.lower()
        return w[:1].upper() + w[1:].lower()
    parts = re.split(r'(\s+)', text)
    out = []
    word_idx = 0
    for p in parts:
        if p.isspace():
            out.append(p)
        else:
            out.append(fix_word(p, is_first=(word_idx == 0)))
            word_idx += 1
    return ''.join(out)

def normalize_whitespace_and_punctuation(text: str) -> str:
    """Insert space before parentheses, collapse spaces."""
    s = text
    # Insert space before (
    s = re.sub(r'(?<=[A-Za-z0-9])\(', ' (', s)
    # Collapse multiple spaces
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def is_abbreviation_of(abbrev: str, full: str) -> bool:
    """
    Check if abbrev is likely an abbreviation of full.
    Returns True if abbrev could be the short form of full.
    
    Examples:
      - "Ach Lia" is abbreviation of "Architectural Liaison"
      - "Pub" is abbreviation of "Publicity"
      - "Sa" is abbreviation of "Security Advisory Section"
    """
    abbrev = abbrev.strip().lower()
    full = full.strip().lower()
    
    if not abbrev or not full or len(abbrev) >= len(full):
        return False
    
    # Split into words
    abbrev_words = abbrev.split()
    full_words = full.split()
    
    # If abbrev has fewer words, check if each abbrev word starts with or is in each full word
    if len(abbrev_words) == len(full_words):
        # Check if each word in abbrev is the first letter(s) of the corresponding full word or appears in it
        for aw, fw in zip(abbrev_words, full_words):
            if not (fw.startswith(aw) or aw in fw):
                return False
        return True
    
    return False


def pick_best_designations(roles_out: list) -> list:
    """
    Remove abbreviations when their full form is present.
    Keeps only the longer/fuller version of designation pairs.
    """
    if len(roles_out) <= 1:
        return roles_out
    
    # Mark roles to remove
    to_remove = set()
    
    for i, role1 in enumerate(roles_out):
        if i in to_remove:
            continue
        for j, role2 in enumerate(roles_out):
            if i == j or j in to_remove:
                continue
            # Check if one is an abbreviation of the other
            if is_abbreviation_of(role1, role2):
                # role1 is abbreviation of role2, so remove role1
                to_remove.add(i)
            elif is_abbreviation_of(role2, role1):
                # role2 is abbreviation of role1, so remove role2
                to_remove.add(j)
    
    return [r for i, r in enumerate(roles_out) if i not in to_remove]

def clean_and_canonicalize_role(raw_role: str) -> str:
    """Clean, expand, canonicalize, and format a single role."""
    if _is_blankish(raw_role):
        return ""
    
    s = str(raw_role).strip()
    
    # Remove TEMP/DES patterns first
    s = re.sub(r'\((?:TEMP|TEMPORARY|DES|DESIGNATE)\)', '', s, flags=re.IGNORECASE)
    
    # Normalize spaces
    s = re.sub(r'\s+', ' ', s).strip()
    
    # Reject pure placeholders
    if not s or s.lower() in {"nan", "none", "null", "-", "()", "", "(temp)", "temp"}:
        return ""
    
    # Apply canonicalization patterns (case-insensitive)
    for pattern, repl in CANON_SYNONYMS.items():
        s = re.sub(pattern, repl, s, flags=re.IGNORECASE)
    
    # Normalize whitespace and punctuation
    s = normalize_whitespace_and_punctuation(s)
    
    # Apply title casing (preserves acronyms)
    s = smart_title_case_role(s)
    
    # Remove repetitive words like "Commander Commander" or "Cdr Cdr"
    s = re.sub(r'\b(\w+)\s+\1(\s+\1)*\b', r'\1', s, flags=re.IGNORECASE)
    
    # Generic letter/number squeeze: "Xyz9" → "Xyz 9"
    s = re.sub(r'([A-Za-z]+)(\d+)', r'\1 \2', s)
    
    # Remove trailing "Team" if it matches the same code
    s = re.sub(r'^(.+\s\d+)\s+Team$', r'\1', s, flags=re.IGNORECASE)
    
    # Final cleanup
    s = re.sub(r'\s+', ' ', s).strip()
    
    # Reject if still ends with (TEMP)
    if re.search(r'\(TEMP(?:ORARY)?\)$', s, flags=re.IGNORECASE):
        return ""
    
    return s

def extract_location_codes_from_row(r, loc_alias):
    """
    Search ALL columns in a row for known location codes.
    Returns the first location code found and its expanded form.
    """
    # Scan all values in the row
    for col_value in r.values:
        if not col_value:
            continue
        text = str(col_value).upper().strip()
        
        # Look for location codes - prioritize longer codes first to avoid partial matches
        for code in sorted(loc_alias.keys(), key=len, reverse=True):
            code_upper = code.upper()
            # Check if code appears as a whole token (with word boundaries)
            if code_upper == text or re.search(r'\b' + re.escape(code_upper) + r'\b', text):
                # Found a location code
                return loc_alias[code]
    
    return ""

def cleanup_role_variants(role):
    """
    Clean up role by:
    1. Replacing 'Rm' with 'Room'
    2. Removing numbered suffixes like (1), (2), (3), etc.
    """
    if not role:
        return role
    
    # Replace Rm with Room
    cleaned = re.sub(r'\bRm\b', 'Room', role, flags=re.IGNORECASE)
    
    # Remove numbered suffixes like (1), (2), etc.
    cleaned = re.sub(r'\s*\(\d+\)\s*$', '', cleaned)
    
    # Normalize whitespace
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    
    return cleaned

def is_abbreviation_of(short, long):
    """Check if short is an abbreviation of long (e.g., 'Ach Lia' for 'Architectural Liaison')."""
    if len(short) >= len(long):
        return False
    short_words = short.upper().split()
    long_words = long.upper().split()
    if len(short_words) > len(long_words):
        return False
    # Check if short words match the first letters of long words or are substrings
    for sw in short_words:
        found = False
        for lw in long_words:
            if lw.startswith(sw) or sw in lw:
                found = True
                break
        if not found:
            return False
    return True

def extract_roles_from_row(r):
    """Extract and canonicalize roles from a row, preferring full forms over abbreviations."""
    
    roles_out = []

    # 1) If Designation Description exists → USE ONLY THAT
    dd_raw = r.get('designation_desc') or ''
    if dd_raw and not is_rank_text(dd_raw):
        dd = clean_and_canonicalize_role(dd_raw)
        if dd:
            roles_out.append(dd)
        # Do NOT read designation at all when desc exists
    else:
        # 2) Otherwise fall back to Designation (try to expand if possible)
        d_raw = r.get('designation') or ''
        if d_raw and not is_rank_text(d_raw):
            d = clean_and_canonicalize_role(d_raw)
            if d:
                roles_out.append(d)

    # 3) Post Type only if non-rank AND only if designation & desc gave nothing
    if not roles_out:
        pt_raw = r.get('post_type') or ''
        if pt_raw and not is_rank_text(pt_raw):
            p = clean_and_canonicalize_role(pt_raw)
            if p:
                roles_out.append(p)

    # Final per-row dedupe
    clean = []
    seen = set()
    for x in roles_out:
        key = x.casefold()
        if key not in seen:
            seen.add(key)
            clean.append(x)

    return clean

def process_excel_file(uploaded_file):
    """Process the uploaded Excel file and return enhanced ranges."""
    try:
        # Read Excel file
        xls = pd.ExcelFile(uploaded_file)
        df = pd.read_excel(uploaded_file, sheet_name=xls.sheet_names[0], header=None)
        
        # Detect header
        header_row = detect_header_row(df)
        df = pd.read_excel(uploaded_file, sheet_name=xls.sheet_names[0], header=header_row)
        df = df.loc[:, ~df.columns.astype(str).str.startswith('Unnamed')]
        df = df.dropna(how='all')
        
        # Normalize columns
        aliases = {
            'date_start': 'date_start', 'date_start_(description)': 'date_start_desc',
            'date_end': 'date_end', 'date_end_(description)': 'date_end_desc',
            'post_type': 'post_type', 'post_type_(description)': 'post_type_desc',
            'designation': 'designation', 'designation_(description)': 'designation_desc',
            'location': 'location', 'location_(description)': 'location_desc',
        }
        df = df.rename(columns={c: aliases.get(snake(c), snake(c)) for c in df.columns})
        
        # Parse dates
        for col in ['date_start', 'date_end']:
            if col in df.columns:
                df[col] = pd.to_datetime(df[col], errors='coerce')
        
        # Ensure columns exist
        for col in ['post_type', 'post_type_desc', 'designation', 'designation_desc', 'location', 'location_desc']:
            if col not in df.columns:
                df[col] = ""
        
        # Sort
        if 'date_start' in df.columns and df['date_start'].notna().any():
            df = df.sort_values(by=['date_start', 'date_end'], ascending=[True, True], na_position='last').reset_index(drop=True)
        
        # Map ranks
        df['reported_rank'] = (df['post_type'].astype(str) + ' || ' + df.get('post_type_desc', '').astype(str)).apply(map_rank)
        df['acting_flag'] = df.apply(is_acting, axis=1)
        
        # True rank calculation
        rows = df.to_dict(orient='records')
        
        def future_has_lower_than(rows, from_idx, ref_rank):
            if ref_rank is None:
                return False
            ref_idx = rank_index.get(ref_rank, -1)
            for j in range(from_idx, len(rows)):
                if bool(rows[j].get('acting_flag')):
                    continue
                rj = rows[j].get('reported_rank')
                if rj is None:
                    continue
                if rank_index.get(rj, -1) < ref_idx:
                    return True
            return False
        
        current_substantive_rank = None
        for i, r in enumerate(rows):
            rep = r.get('reported_rank')
            acting = bool(r.get('acting_flag'))
            
            if current_substantive_rank is None:
                if rep is None:
                    r['true_rank'] = None
                else:
                    looks_acting = acting or future_has_lower_than(rows, i + 1, rep)
                    if looks_acting:
                        r['true_rank'] = None
                    else:
                        current_substantive_rank = rep
                        r['true_rank'] = current_substantive_rank
                continue
            
            if rep is None:
                r['true_rank'] = current_substantive_rank
                continue
            
            if acting:
                r['true_rank'] = current_substantive_rank
                continue
            
            cur_idx = rank_index.get(current_substantive_rank, -1)
            rep_idx = rank_index.get(rep, -1)
            
            if rep_idx > cur_idx:
                if future_has_lower_than(rows, i + 1, rep):
                    r['true_rank'] = current_substantive_rank
                else:
                    current_substantive_rank = rep
                    r['true_rank'] = current_substantive_rank
            else:
                r['true_rank'] = current_substantive_rank
        
        first_true = next((rr.get('true_rank') for rr in rows if rr.get('true_rank') is not None), None)
        if first_true is not None:
            for r in rows:
                if r.get('true_rank') is None:
                    r['true_rank'] = first_true
        
        df = pd.DataFrame(rows)
        
        # Year ranges by rank
        segments = []
        current_rank = None
        seg_start = None
        seg_end = None
        
        def max_dt(a, b):
            if pd.isna(a): return b
            if pd.isna(b): return a
            return max(a, b)
        
        def min_dt(a, b):
            if pd.isna(a): return b
            if pd.isna(b): return a
            return min(a, b)
        
        for _, row in df.iterrows():
            tr = row.get('true_rank')
            ds = row.get('date_start')
            de = row.get('date_end')
            if current_rank is None:
                current_rank = tr
                seg_start = ds
                seg_end = de
                continue
            if tr != current_rank:
                segments.append({'true_rank': current_rank, 'start': seg_start, 'end': seg_end})
                current_rank = tr
                seg_start = ds
                seg_end = de
            else:
                seg_start = min_dt(seg_start, ds)
                seg_end = max_dt(seg_end, de)
        
        if current_rank is not None:
            segments.append({'true_rank': current_rank, 'start': seg_start, 'end': seg_end})
        
        def fmt_year_range(start_dt, end_dt):
            if pd.isna(start_dt) and pd.isna(end_dt):
                return "Unknown"
            if pd.isna(start_dt):
                return f"–{end_dt.year}"
            sy = start_dt.year
            if pd.isna(end_dt):
                return f"{sy}–Present"
            ey = end_dt.year
            if ey < sy:
                return f"{sy}"
            return f"{sy}–{ey}"
        
        year_ranges = [
            {'true_rank': seg['true_rank'], 'start': seg.get('start'), 'end': seg.get('end'),
             'year_range': fmt_year_range(seg.get('start'), seg.get('end'))}
            for seg in segments if seg.get('true_rank') is not None
        ]
        
        # Enhanced ranges with locations and roles
        enhanced_ranges = []
        for seg in year_ranges:
            tr = seg['true_rank']
            start_dt = seg['start']
            end_dt = seg['end']
            
            mask = (df['true_rank'] == tr)
            if pd.notna(start_dt):
                mask &= (df['date_start'].isna() | (df['date_start'] >= start_dt))
            if pd.notna(end_dt):
                mask &= (df['date_end'].isna() | (df['date_end'] <= end_dt))
            
            sub = df.loc[mask].copy()
            
            loc_desc = sub['location_desc']
            loc_series = loc_desc.where(
                loc_desc.notna() & (loc_desc.astype(str).str.strip() != ''),
                sub['location']
            ).apply(lambda x: normalize_location(x, STARTER_LOCATION_ALIASES))
            
            # Also search each row for location codes in ANY column
            row_location_codes = sub.apply(lambda row: extract_location_codes_from_row(row, STARTER_LOCATION_ALIASES), axis=1)
            
            # Prioritize location codes found in row over division column
            loc_series = loc_series.where(row_location_codes == '', row_location_codes)
            
            sub['role_list'] = sub.apply(extract_roles_from_row, axis=1)
            
            roles_by_loc = {}
            seen_by_loc = {}
            
            for l, roles_here in zip(loc_series, sub['role_list']):
                if not l:
                    continue
                roles_by_loc.setdefault(l, [])
                seen_by_loc.setdefault(l, set())
                for role in roles_here:
                    if not role or role.upper() == "LEAVE RESERVE":
                        continue
                    # Canonicalize role one more time for location-level dedup
                    role_canonical = clean_and_canonicalize_role(role)
                    if not role_canonical:
                        continue
                    key = role_canonical.casefold()
                    if key not in seen_by_loc[l]:
                        roles_by_loc[l].append(role_canonical)
                        seen_by_loc[l].add(key)
            
            # Remove abbreviations from roles (e.g., "Ach Lia" if "Architectural Liaison" exists)
            for loc in roles_by_loc:
                roles_list = roles_by_loc[loc]
                roles_to_remove = set()
                for i, role1 in enumerate(roles_list):
                    for j, role2 in enumerate(roles_list):
                        if i != j:
                            # Check both directions - role1 could be abbrev of role2 OR vice versa
                            if is_abbreviation_of(role1, role2):
                                roles_to_remove.add(role1)
                            elif is_abbreviation_of(role2, role1):
                                roles_to_remove.add(role2)
                # Remove abbreviations, preserving order
                roles_by_loc[loc] = [r for r in roles_list if r not in roles_to_remove]
            
            unique_locs = []
            seen_locs = set()
            for l in loc_series.tolist():
                if not l:
                    continue
                if l not in seen_locs:
                    seen_locs.add(l)
                    unique_locs.append(l)
            
            # ===== DIVISION → DISTRICT MERGE =====
            # Recognize DIV/DIVISION and DIST/DISTRICT patterns
            def extract_base_and_type(loc_name: str):
                """Extract base and type from location name."""
                s = str(loc_name or "").strip().upper()
                s = re.sub(r'\s+', ' ', s)
                
                typ = None
                if re.search(r'\bDIV(?:ISION)?\.?\b', s):
                    typ = 'DIVISION'
                elif re.search(r'\bDIST(?:RICT)?\.?\b', s):
                    typ = 'DISTRICT'
                
                # Remove type tokens to get base
                base = s
                base = re.sub(r'\bDIV(?:ISION)?\.?\b', '', base)
                base = re.sub(r'\bDIST(?:RICT)?\.?\b', '', base)
                base = re.sub(r'\s+', ' ', base).strip()
                
                return base, typ
            
            # Map each base to its locations and types
            base_to_locs = {}
            for loc_name in unique_locs:
                base, typ = extract_base_and_type(loc_name)
                if base:
                    base_to_locs.setdefault(base, {'DIVISION': None, 'DISTRICT': None})
                    if typ:
                        base_to_locs[base][typ] = loc_name
            
            # Find Division→District merges and collect them
            division_to_district_merge = {}
            for base, locs_dict in base_to_locs.items():
                if locs_dict['DIVISION'] and locs_dict['DISTRICT']:
                    # Both exist: merge Division into District
                    div_loc = locs_dict['DIVISION']
                    dist_loc = locs_dict['DISTRICT']
                    division_to_district_merge[div_loc] = dist_loc
            
            # Merge roles from Divisions into Districts
            for div_loc, dist_loc in division_to_district_merge.items():
                div_roles = roles_by_loc.pop(div_loc, [])
                if dist_loc in roles_by_loc:
                    # Merge: add div roles to district, deduplicate
                    seen_in_dist = {r.casefold() for r in roles_by_loc[dist_loc]}
                    for r in div_roles:
                        if r.casefold() not in seen_in_dist:
                            roles_by_loc[dist_loc].append(r)
                            seen_in_dist.add(r.casefold())
            
            # Clean up all roles: standardize Room/Rm, remove numbered variants, deduplicate
            for loc_name in roles_by_loc:
                roles = roles_by_loc[loc_name]
                # Canonicalize all roles again to ensure all variants are expanded
                canonical_roles = [clean_and_canonicalize_role(r) for r in roles]
                # Deduplicate after canonicalization
                unique_roles = []
                seen = set()
                for r in canonical_roles:
                    key = re.sub(r'[^a-z0-9]', '', r.casefold())
                    if key not in seen and r:
                        seen.add(key)
                        unique_roles.append(r)
                # Fuzzy deduplication: remove any role that is a substring, abbreviation, or fuzzy match of a longer role in the same list
                to_remove = set()
                def norm(s):
                    return re.sub(r'[^a-z0-9]', '', s.casefold())
                for i, role1 in enumerate(unique_roles):
                    for j, role2 in enumerate(unique_roles):
                        if i == j:
                            continue
                        n1, n2 = norm(role1), norm(role2)
                        # Remove role1 if it is a substring or abbreviation of role2
                        if n1 != n2 and (n1 in n2 or is_abbreviation_of(role1, role2)):
                            to_remove.add(i)
                roles_by_loc[loc_name] = [r for i, r in enumerate(unique_roles) if i not in to_remove]
            
            # Rebuild unique locations list without Divisions
            final_unique_locs = [loc for loc in unique_locs if loc not in division_to_district_merge]
            
            enhanced_ranges.append({
                'true_rank': tr,
                'year_range': seg['year_range'],
                'locations': final_unique_locs,
                'roles_by_location': roles_by_loc
            })
        
        return enhanced_ranges, None
    
    except Exception as e:
        return None, str(e)

def generate_word_document(enhanced_ranges):
    """Generate Word document with a table format."""
    doc = Document()
    
    # Rank expansion mapping with full, official rank names
    rank_full_names = {
        'PC': 'Police Constable',
        'SPC': 'Senior Police Constable',
        'SGT': 'Sergeant',
        'SSGT': 'Station Sergeant',
        'PI': 'Probationary Inspector',
        'IP': 'Inspector',
        'SIP': 'Senior Inspector',
        'IP/SIP': 'Inspector / Senior Inspector',
        'CIP': 'Chief Inspector',
        'SP': 'Superintendent',
        'SSP': 'Senior Superintendent',
        'CSP': 'Chief Superintendent',
        'ACP': 'Assistant Commissioner',
        'SACP': 'Senior Assistant Commissioner',
        'DCP': 'Deputy Commissioner',
        'CP': 'Commissioner',
    }
    
    # Helper function to shade cell
    def shade_cell(cell, color):
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Create table with 3 columns (Year Range, Rank, Posting Info)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Set column widths (narrow for Year Range and Rank, wide for Posting Info)
    table.columns[0].width = Inches(1.0)  # Year Range - narrow
    table.columns[1].width = Inches(1.2)  # Rank - narrow
    table.columns[2].width = Inches(4.3)  # Posting Location & Roles - wide
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Year Range"
    header_cells[1].text = "Rank"
    header_cells[2].text = "Posting Location & Roles"
    
    # Format header row - bold, size 13, light blue background
    for cell in header_cells:
        shade_cell(cell, "ADD8E6")  # Light blue
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(13)
    
    # Add data rows
    for item in enhanced_ranges:
        year_range = item['year_range']
        rank = item['true_rank']
        locations = item['locations']
        roles_by_location = item['roles_by_location']
        
        # Convert rank acronym to full name
        rank_display = rank_full_names.get(rank, rank)
        
        # Build posting info text
        posting_info = []
        if locations:
            for i, loc in enumerate(locations):
                # Add blank line before location (except for first one)
                if i > 0:
                    posting_info.append("")
                posting_info.append(f"{loc}")
                roles = roles_by_location.get(loc, [])
                if roles:
                    for role in roles:
                        posting_info.append(f"  • {role}")
        else:
            posting_info.append("(No locations recorded)")
        
        posting_text = "\n".join(posting_info)
        
        # Add row
        row_cells = table.add_row().cells
        row_cells[0].text = year_range
        row_cells[1].text = rank_display
        row_cells[2].text = posting_text
        
        # Format data rows - white background, size 13
        for cell in row_cells:
            shade_cell(cell, "FFFFFF")  # White
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(13)
    
    return doc

# ========= STREAMLIT UI =========
st.markdown("---")

# File uploader
col1, col2 = st.columns([2, 1])
with col1:
    uploaded_file = st.file_uploader("📁 Upload Excel File", type=["xlsx", "xls"])

if uploaded_file is not None:
    st.success(f"✓ File uploaded: {uploaded_file.name}")
    
    # Process button
    if st.button("🔄 Process File", use_container_width=True):
        with st.spinner("Processing your Excel file..."):
            enhanced_ranges, error = process_excel_file(uploaded_file)
            
            if error:
                st.error(f"Error processing file: {error}")
            else:
                st.success("✓ Processing complete!")
                
                # Generate Word document
                doc = generate_word_document(enhanced_ranges)
                
                # Create downloadable file
                doc_bytes = io.BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                # Download button
                st.download_button(
                    label="📥 Download Word Document",
                    data=doc_bytes.getvalue(),
                    file_name="HKPF_Posting_Summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                # Show preview
                with st.expander("📋 Preview Summary"):
                    for item in enhanced_ranges:
                        st.write(f"**{item['true_rank']}: {item['year_range']}**")
                        for loc in item['locations']:
                            st.write(f"  • {loc}")
                            roles = item['roles_by_location'].get(loc, [])
                            for role in roles:
                                st.write(f"    - {role}")
else:
    st.info("👆 Please upload an Excel file to get started")
